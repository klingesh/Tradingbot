"""
Generates docs/TRADING_BOT_FULL_DOCUMENTATION.docx - the complete build log:
every phase, command run, error encountered, strategy tested, and result.

Usage:
    pip install python-docx
    python docs/generate_full_documentation.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

OUT = os.path.join(os.path.dirname(__file__), "TRADING_BOT_FULL_DOCUMENTATION.docx")

MONO = "Consolas"
CODE_BG = RGBColor(0x1E, 0x1E, 0x1E)
CODE_FG = RGBColor(0x2E, 0x2E, 0x2E)
ERR_FG = RGBColor(0xB0, 0x30, 0x20)
OK_FG = RGBColor(0x1E, 0x7A, 0x3C)


def h1(d, t):
    d.add_heading(t, level=1)


def h2(d, t):
    d.add_heading(t, level=2)


def h3(d, t):
    d.add_heading(t, level=3)


def para(d, t, bold=False, italic=False, size=11, color=None):
    p = d.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def bullets(d, items, style="List Bullet"):
    for it in items:
        d.add_paragraph(str(it), style=style)


def code(d, text, lang="", fg=CODE_FG):
    if lang:
        cp = d.add_paragraph()
        cr = cp.add_run(f"[{lang}]")
        cr.bold = True
        cr.font.size = Pt(8)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x88)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(8.5)
    r.font.color.rgb = fg
    return p


def error_block(d, text):
    code(d, text, lang="error output", fg=ERR_FG)


def table(d, headers, rows):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = str(htxt)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    d.add_paragraph()
    return t


def build() -> None:
    d = Document()

    # ---------------- TITLE ----------------
    t = d.add_heading("Multi-Asset Algorithmic Trading Bot", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = para(d, "Complete Development Documentation — Every Step, Command, Error and Strategy",
             italic=True, size=13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = para(d, "Repository: github.com/klingesh/Tradingbot   |   Platform: MetaTrader 5 (JustMarkets)",
             size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = para(d, "Language: Python 3.11–3.14   |   Status: Deployed 24/7 on DEMO account", size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d.add_paragraph()
    para(d, "DISCLAIMER", bold=True, color=ERR_FG)
    para(d, "This project is for research and education only. Trading carries substantial "
            "risk of loss. Nothing in this document is financial advice. All performance "
            "figures are BACKTESTS / out-of-sample simulations, not live trading returns. "
            "Past performance does not predict future results. The system is running on a "
            "DEMO account; only risk capital you can afford to lose.", size=9, italic=True)

    d.add_page_break()

    # ---------------- CONTENTS ----------------
    h1(d, "Contents")
    bullets(d, [
        "1. Project Goal and the Critical Reframing",
        "2. Technology Stack and Environment",
        "3. Phase 1 — Foundations: Risk Engine",
        "4. Phase 2 — Data Pipeline (and the geo-block errors)",
        "5. Phase 3 — First Strategy and the Honest Backtester",
        "6. Phase 4 — Edge Hunting: the Higher-Timeframe Filter",
        "7. Phase 5 — Strategy Archetype Comparison",
        "8. Phase 6 — Walk-Forward Validation (anti-overfitting)",
        "9. Phase 7 — Multi-Asset Expansion (forex & commodities)",
        "10. Phase 8 — Mean-Reversion Discovery",
        "11. Phase 9 — News / Economic Calendar Layer",
        "12. Phase 10 — MT5 Connector and Live Bot",
        "13. Phase 11 — Windows Setup: every error we hit",
        "14. Phase 12 — Performance Logging and VPS 24/7",
        "15. Phase 13 — Mining 344 SMC/ICT Indicators",
        "16. Phase 14 — Volatility Targeting (evidence-based)",
        "17. Phase 15 — Universe Expansion & Diversification",
        "18. Phase 16 — News Awareness Tools",
        "19. Phase 17 — Consolidation and Config Decoupling",
        "20. Complete File Inventory",
        "21. All Strategies: Logic and Verdicts",
        "22. Complete Error Log and Resolutions",
        "23. Final Portfolio and Results",
        "24. Key Lessons Learned",
        "25. Operating Manual (commands reference)",
    ], style="List Number")

    d.add_page_break()

    # ---------------- 1. GOAL ----------------
    h1(d, "1. Project Goal and the Critical Reframing")
    h2(d, "Original request")
    para(d, "Build a hedging and trading bot for forex, crypto and commodities that:")
    bullets(d, [
        "Trades autonomously (assigns lot size, risk %, take-profit, stop-loss)",
        "Takes trades on news impacts and government data releases",
        "Combines fundamental and technical analysis",
        "Targets a 70–80% win rate",
    ])

    h2(d, "The reframing that shaped everything")
    para(d, "The 70–80% win-rate target was challenged immediately, because win rate is "
            "one of the most misleading metrics in trading. A bot can win 80% of trades "
            "and still lose money if losses are large; a bot can win 40% and be highly "
            "profitable if winners are bigger.")
    code(d, "Expectancy = (Win% x AvgWin) - (Loss% x AvgLoss)\n\n"
            "A common way to FAKE a high win rate:\n"
            "  tiny take-profits + huge/no stop-loss\n"
            "  -> wins often, then one trade erases 30 wins", lang="the core formula")
    para(d, "New objective: POSITIVE EXPECTANCY WITH CONTROLLED DRAWDOWN, validated "
            "out-of-sample. Win rate became a side effect, not a target.", bold=True)
    para(d, "This was later confirmed by our own data: the final profitable strategies win "
            "only ~40–57% of the time, and the highest-win-rate strategy we tested "
            "(mean-reversion on crypto, 45%) was the WORST performer (-32%).")

    # ---------------- 2. STACK ----------------
    h1(d, "2. Technology Stack and Environment")
    table(d, ["Component", "Choice", "Reason"], [
        ["Language", "Python 3.11–3.14", "Best ecosystem; official MT5 package"],
        ["Broker/platform", "MetaTrader 5 (JustMarkets)", "User's existing account; multi-asset in one API"],
        ["MT5 bridge", "MetaTrader5 pip package", "Official; WINDOWS ONLY"],
        ["Data (crypto)", "OKX public API", "Free, no auth, long history"],
        ["Data (FX/commodities/indices)", "Yahoo Finance chart API", "Free, no auth, 15y daily / 2y hourly"],
        ["Economic calendar", "ForexFactory community JSON", "Free high-impact events + forecast/actual"],
        ["News headlines", "Google News RSS", "Free, no key, dated headlines"],
        ["Analysis", "pandas, numpy", "Indicators hand-written (no TA lib dependency)"],
        ["Config", "PyYAML", "Human-editable settings"],
        ["Testing", "pytest", "31 automated tests"],
        ["Deployment", "Windows VPS + .bat auto-restart", "24/7 without the laptop"],
        ["Version control", "Git / GitHub", "klingesh/Tradingbot"],
    ])
    para(d, "Critical constraint discovered early:", bold=True)
    para(d, "The MetaTrader5 Python package only runs on Windows. Therefore the codebase "
            "was designed so all logic (risk, strategies, backtesting, news) is pure "
            "Python and testable anywhere, while the MT5 connector is import-guarded and "
            "runs only on the user's Windows machine/VPS.")

    # ---------------- 3. RISK ENGINE ----------------
    h1(d, "3. Phase 1 — Foundations: Risk Engine")
    para(d, "Built FIRST, before any strategy — it is the survival layer.")
    h2(d, "File created: src/risk/position_sizing.py")
    para(d, "Converts a risk percentage into a broker-legal lot size, using the stop "
            "distance and the instrument's real contract specs. Currency-agnostic so it "
            "works identically on cent and standard accounts.")
    code(d, "risk_amount        = balance * (risk_percent / 100)\n"
            "ticks_to_stop      = stop_loss_distance_price / tick_size\n"
            "money_risk_per_lot = ticks_to_stop * tick_value\n"
            "raw_lots           = risk_amount / money_risk_per_lot\n"
            "lots               = round_DOWN_to_step(raw_lots)   # never over-risk",
         lang="python — core sizing formula")
    para(d, "Safety features built in:")
    bullets(d, [
        "Hard cap on risk % per trade (never exceed max, even if configured higher)",
        "Rounds lot size DOWN to the broker's volume step (never rounds risk up)",
        "MinLotPolicy.SKIP — if even the minimum lot risks too much, skip the trade",
        "Rejects zero/negative balance and zero stop distance",
    ])
    h2(d, "Command run and result")
    code(d, "python -m pytest tests/ -v", lang="bash")
    code(d, "8 passed in 0.02s", lang="output", fg=OK_FG)
    para(d, "Also created: config/config.yaml (risk settings), tests/test_position_sizing.py, "
            "README.md, requirements.txt.")

    # ---------------- 4. DATA ----------------
    h1(d, "4. Phase 2 — Data Pipeline (and the geo-block errors)")
    h2(d, "ERROR 1 — Binance geo-blocked")
    code(d, 'urllib.request.urlopen("https://api.binance.com/api/v3/klines?...")', lang="python")
    error_block(d, "Binance failed: <HTTPError 451: ''>\n"
                   "(451 = Unavailable For Legal Reasons — geo-restricted)")
    h2(d, "Resolution: tested five alternatives")
    code(d, "XX  CryptoCompare : HTTPError 401 Unauthorized\n"
            "OK  Kraken        : reachable\n"
            "XX  Coinbase      : HTTPError 400 Bad Request\n"
            "OK  OKX           : reachable      <-- CHOSEN (best pagination/history)\n"
            "XX  Bybit         : HTTPError 403 Forbidden", lang="probe results")
    para(d, "Created src/data/loader.py — OKX loader with CSV caching and pagination.")
    code(d, "Total bars: 6000\nRange: 2023-10-12 -> 2026-07-08\nApprox years: 2.74",
         lang="output", fg=OK_FG)

    h2(d, "ERROR 2 — Yahoo silently returned MONTHLY bars")
    para(d, "When forex/commodities were added later, using range=max returned only ~270 "
            "rows for 20+ years — Yahoo downsamples huge ranges to monthly.")
    error_block(d, "EURUSD (forex) 273 bars 2003-12-01 -> 2026-07-10   <-- monthly, not daily!")
    para(d, "Resolution: use explicit period1/period2 unix timestamps instead of range=max.")
    code(d, "period2 = int(time.time())\n"
            "period1 = period2 - int(years * 365.25 * 24 * 3600)\n"
            "url += f'?interval={interval}&period1={period1}&period2={period2}'", lang="python fix")
    code(d, "EURUSD (forex) 3906 bars 2011-07-10 -> 2026-07-10   <-- true daily",
         lang="output", fg=OK_FG)

    h2(d, "ERROR 3 — Yahoo hourly range limit")
    error_block(d, "EURUSD 1h 2.5 years -> <HTTPError 422: 'Unprocessable Entity'>")
    para(d, "Resolution: Yahoo caps hourly data at ~729 days. Fetch 729 days of 1h bars "
            "and resample to H4 in pandas, giving ~3,100 H4 bars per instrument.")
    code(d, 'h4 = hourly.resample("4h").agg({"open":"first","high":"max",\n'
            '                                "low":"min","close":"last","volume":"sum"})\n'
            '       .dropna(subset=["open","high","low","close"])', lang="python")

    # ---------------- 5. BACKTESTER ----------------
    h1(d, "5. Phase 3 — First Strategy and the Honest Backtester")
    h2(d, "Files created")
    bullets(d, [
        "src/strategy/indicators.py — EMA, RSI (Wilder), ATR, hand-written in pandas",
        "src/strategy/base.py — abstract Strategy class (signals only; no risk logic)",
        "src/strategy/ema_rsi_swing.py — EMA crossover + RSI confirmation + ATR stops",
        "src/backtest/metrics.py — expectancy, PF, Sharpe, max drawdown, R-multiples",
        "src/backtest/engine.py — event-driven backtester",
    ])
    h2(d, "Rules that keep the backtest honest (no cheating)")
    bullets(d, [
        "Signals decided on a bar's CLOSE; entry at the NEXT bar's OPEN",
        "SL/TP checked intrabar using each bar's high/low",
        "If SL and TP could both hit in one bar, assume SL first (worst case)",
        "Spread/slippage cost applied to EVERY fill (adverse direction)",
        "Position size from the real risk engine, sized off current equity",
        "One position at a time (swing style)",
    ])
    h2(d, "First result — the honest disappointment")
    code(d, "python backtest_run.py", lang="bash")
    table(d, ["Metric", "Value", "Meaning"], [
        ["Trades", "107", "decent sample"],
        ["Win rate", "32.7%", "low — but not the point"],
        ["Avg win / loss", "+1.67R / -0.81R", "winners ~2x losers"],
        ["Profit factor", "0.98", "<1 = not profitable"],
        ["Expectancy", "~0R", "BREAK-EVEN"],
        ["Return", "-2.38%", "slightly negative over 2.7 years"],
        ["Max drawdown", "18.6%", ""],
    ])
    para(d, "Lesson: the first strategy had NO edge. That is normal and the backtester's "
            "job is to say so. Also note it won only 33% of trades yet was near "
            "break-even because winners were 2x losers — the win-rate lesson made concrete.")
    h2(d, "Bug fixed: misleading verdict label")
    para(d, "The report labelled +0.003R as 'POSITIVE EDGE'. Corrected the thresholds:")
    code(d, "if self.profit_factor > 1.05 and self.expectancy_r > 0.03:\n"
            "    edge = 'POSITIVE EDGE'\n"
            "elif self.profit_factor >= 0.97:\n"
            "    edge = 'BREAK-EVEN (no real edge)'\n"
            "else:\n"
            "    edge = 'NEGATIVE (losing)'", lang="python fix")

    # ---------------- 6. HTF FILTER ----------------
    h1(d, "6. Phase 4 — Edge Hunting: the Higher-Timeframe Filter")
    para(d, "Hypothesis: only take H4 trades aligned with the DAILY trend (don't fight "
            "the bigger picture).")
    h2(d, "Anti-look-ahead detail (critical)")
    para(d, "A daily bar's trend is only known at that day's close, so the higher-timeframe "
            "series is shifted by one bar before being mapped onto intraday bars.")
    code(d, "htf_close = out['close'].resample(self.htf_rule).last().dropna()\n"
            "htf_ema   = ema(htf_close, self.htf_ema_period)\n"
            "htf_dir   = (htf_close > htf_ema).map({True: 1, False: -1})\n"
            "htf_dir   = htf_dir.shift(1)           # only completed HTF bar\n"
            "return htf_dir.reindex(out.index, method='ffill')", lang="python")
    h2(d, "Result — a real edge appeared")
    table(d, ["Variant", "Trades", "Win", "PF", "Exp(R)", "Return", "MaxDD", "Sharpe"], [
        ["Baseline (no filter)", "107", "32.7%", "0.98", "+0.003R", "-0.50%", "18.61%", "0.06"],
        ["+ Daily EMA50 filter", "41", "41.5%", "1.37", "+0.239R", "+19.32%", "11.11%", "0.65"],
        ["+ Daily EMA100 filter", "41", "41.5%", "1.35", "+0.231R", "+18.70%", "8.80%", "0.63"],
    ])
    para(d, "Fewer, better trades (107 -> 41); profit factor crossed 1.0; drawdown nearly "
            "halved. Win rate still only 41% — profitable because winners are bigger.")

    # ---------------- 7. ARCHETYPES ----------------
    h1(d, "7. Phase 5 — Strategy Archetype Comparison")
    para(d, "Created src/strategy/breakout.py (Donchian channel) and "
            "src/strategy/mean_reversion.py (Bollinger + RSI), then compared all on BTC H4.")
    table(d, ["Strategy", "Trades", "Win", "PF", "Exp(R)", "Return", "MaxDD", "Sharpe"], [
        ["EMA+RSI (no filter)", "107", "32.7%", "0.98", "+0.003R", "-0.50%", "18.6%", "0.06"],
        ["EMA+RSI + daily filter", "41", "41.5%", "1.37", "+0.239R", "+19.3%", "11.1%", "0.65"],
        ["Donchian breakout", "206", "35.0%", "1.08", "+0.069R", "+26.7%", "28.9%", "0.49"],
        ["Donchian + daily filter", "116", "37.1%", "1.12", "+0.100R", "+20.1%", "21.7%", "0.45"],
        ["Bollinger mean-reversion", "162", "45.1%", "0.78", "-0.111R", "-32.3%", "32.3%", "-0.80"],
    ])
    para(d, "THE KEY OBSERVATION: mean-reversion had the HIGHEST win rate (45%) and was "
            "the WORST performer (-32%). Meanwhile breakout made the highest raw return "
            "(+26.7%) but with a brutal 29% drawdown. Winner on a risk-adjusted basis: "
            "EMA+RSI + daily filter.", bold=True)

    # ---------------- 8. WALK FORWARD ----------------
    h1(d, "8. Phase 6 — Walk-Forward Validation (anti-overfitting)")
    para(d, "The most important test in the project. Procedure:")
    bullets(d, [
        "Take an in-sample (IS) window of past bars",
        "Grid-search parameters on it",
        "Apply the chosen parameters to the NEXT out-of-sample (OOS) window",
        "Roll windows forward, repeat, stitch all OOS segments together",
    ], style="List Number")
    para(d, "Engine change required: allow trading to be restricted to a window while "
            "indicators still compute on full history.")
    code(d, "ts = data[ts_col]\n"
            "lo = 1 if trade_start is None else max(1, int((ts < trade_start).sum()))\n"
            "hi = len(data) if trade_end is None else int((ts < trade_end).sum())\n"
            "for i in range(lo, hi):   # trades confined to the OOS window", lang="python")
    h2(d, "Result — the edge SURVIVED out-of-sample")
    code(d, "python experiments/walk_forward_test.py", lang="bash")
    table(d, ["Metric", "Aggregate OOS value"], [
        ["Trades", "43"],
        ["Win rate", "41.86%"],
        ["Profit factor", "1.39"],
        ["Avg win / loss", "+2.04R / -1.02R"],
        ["Expectancy", "+0.264R"],
        ["Total return", "+22.79%"],
        ["Max drawdown", "12.65%"],
        ["Sharpe", "0.84"],
    ])
    para(d, "7 of 9 OOS folds were profitable. Because the edge held on unseen data AND "
            "came from a sound principle (trade with the higher-timeframe trend), it was "
            "judged genuine rather than curve-fitted.")

    # ---------------- 9. MULTI ASSET ----------------
    h1(d, "9. Phase 7 — Multi-Asset Expansion (forex & commodities)")
    para(d, "User request: avoid crypto due to volatility; test forex and commodities.")
    h2(d, "PERFORMANCE PROBLEM — backtest too slow")
    error_block(d, "Command timed out after 120000ms\n"
                   "(row-by-row pandas .iloc loop; 2,500+ backtests required)")
    para(d, "Resolution: rewrote the engine hot loop with NumPy arrays (~20x faster). "
            "Verified identical results via regression check.")
    code(d, "ts_index = signals.index.to_numpy()\n"
            "open_ = signals['open'].to_numpy(dtype=float)\n"
            "high  = signals['high'].to_numpy(dtype=float)\n"
            "low   = signals['low'].to_numpy(dtype=float)\n"
            "close = signals['close'].to_numpy(dtype=float)\n"
            "sig   = signals['signal'].to_numpy()", lang="python optimization")
    code(d, "Regression check: 107 trades, PF 0.98 — identical to before", lang="output", fg=OK_FG)
    h2(d, "Daily-timeframe results (15 years, weekly trend filter)")
    table(d, ["Instrument", "Class", "PF", "Exp(R)", "Return", "MaxDD", "Sharpe"], [
        ["EURUSD", "forex", "0.63", "-0.237", "-6.10%", "11.26%", "-0.19"],
        ["GBPUSD", "forex", "1.45", "+0.254", "+11.45%", "8.27%", "0.27"],
        ["USDJPY", "forex", "2.76", "+0.619", "+27.63%", "8.17%", "0.58"],
        ["AUDUSD", "forex", "1.28", "+0.167", "+10.26%", "7.33%", "0.24"],
        ["USDCAD", "forex", "0.65", "-0.237", "-6.31%", "11.67%", "-0.21"],
        ["XAUUSD", "commodity", "0.94", "-0.022", "-1.46%", "15.19%", "-0.01"],
        ["XAGUSD", "commodity", "0.71", "-0.179", "-6.29%", "14.27%", "-0.13"],
        ["WTI", "commodity", "2.53", "+0.530", "+24.29%", "3.90%", "0.59"],
    ])
    para(d, "Finding: trend-following is MIXED on forex/commodities — strong where markets "
            "trended (USDJPY, WTI) but losing on range-bound pairs (EURUSD, USDCAD). "
            "Average Sharpe only 0.14 vs crypto's 0.84, and very few trades (13–23 in 12 years).")
    h2(d, "H4 timeframe results — trend FAILS on forex majors")
    table(d, ["Instrument", "PF", "Return", "MaxDD", "Sharpe"], [
        ["EURUSD", "0.55", "-13.59%", "21.27%", "-0.86"],
        ["GBPUSD", "0.60", "-9.42%", "18.71%", "-0.48"],
        ["USDJPY", "0.58", "-8.44%", "15.41%", "-0.57"],
        ["AUDUSD", "0.61", "-12.46%", "13.36%", "-0.83"],
        ["USDCAD", "1.27", "+4.03%", "8.59%", "0.32"],
        ["XAUUSD (gold)", "1.87", "+13.90%", "5.95%", "1.35"],
        ["XAGUSD (silver)", "1.61", "+11.02%", "6.94%", "0.85"],
        ["WTI", "1.09", "+1.66%", "10.65%", "0.16"],
    ])
    para(d, "Forex majors whipsaw intraday and fail trend-following at H4, while "
            "commodities (gold Sharpe 1.35) trend beautifully. This motivated testing "
            "mean-reversion on the majors.")

    # ---------------- 10. MEAN REVERSION ----------------
    h1(d, "10. Phase 8 — Mean-Reversion Discovery")
    para(d, "Head-to-head on forex majors at H4: trend vs mean-reversion.")
    table(d, ["Pair", "Strategy", "PF", "Exp(R)", "Return", "MaxDD", "Sharpe"], [
        ["EURUSD", "trend", "0.55", "-0.339", "-13.59%", "21.27%", "-0.86"],
        ["EURUSD", "mean-revert", "0.74", "-0.140", "-6.69%", "11.35%", "-0.50"],
        ["GBPUSD", "trend", "0.60", "-0.289", "-9.42%", "18.71%", "-0.48"],
        ["GBPUSD", "mean-revert", "0.86", "-0.074", "-6.11%", "13.53%", "-0.40"],
        ["USDJPY", "trend", "0.58", "-0.292", "-8.44%", "15.41%", "-0.57"],
        ["USDJPY", "mean-revert", "1.76", "+0.373", "+17.88%", "11.94%", "1.18"],
        ["AUDUSD", "trend", "0.61", "-0.264", "-12.46%", "13.36%", "-0.83"],
        ["AUDUSD", "mean-revert", "1.60", "+0.306", "+18.87%", "9.09%", "1.27"],
        ["USDCAD", "trend", "1.27", "+0.187", "+4.03%", "8.59%", "0.32"],
        ["USDCAD", "mean-revert", "0.43", "-0.430", "-21.62%", "23.96%", "-1.49"],
    ])
    para(d, "CORE INSIGHT OF THE PROJECT: match the strategy to the instrument's "
            "character. USDJPY went from -8% (trend) to +18% (mean-reversion) — the same "
            "market, opposite outcome, purely from using the right tool.", bold=True)

    # ---------------- 11. NEWS ----------------
    h1(d, "11. Phase 9 — News / Economic Calendar Layer")
    h2(d, "Data source testing")
    code(d, "XX  FMP economic calendar (demo key) : HTTPError 401 Unauthorized\n"
            "OK  ForexFactory community JSON      : 98 events, 14 High impact  <-- CHOSEN",
         lang="probe results")
    para(d, "The feed includes impact level, forecast and previous values — enabling a "
            "future 'surprise' module.")
    h2(d, "Files created")
    bullets(d, [
        "src/news/calendar.py — DeterministicUSCalendar (NFP = first Friday, FOMC "
        "published dates, DST-correct via zoneinfo) for BACKTESTS; ForexFactoryCalendar "
        "for LIVE",
        "src/news/filter.py — NewsFilter with a vectorized blackout mask",
    ])
    h2(d, "BUG 1 — timezone comparison crash")
    error_block(d, "TypeError: can't compare offset-naive and offset-aware datetimes\n"
                   "  in np.searchsorted(lo_sorted, bar_times, side='right')")
    code(d, "# Fix: normalize everything to naive-UTC datetime64\n"
            "self._event_times = np.array(sorted(\n"
            "    e.time_utc.tz_convert(UTC).tz_localize(None).to_datetime64()\n"
            "    for e in events), dtype='datetime64[ns]')", lang="python fix")
    h2(d, "BUG 2 — bar duration ignored")
    para(d, "Only the bar's OPEN time was checked, so an event landing mid-bar (e.g. FOMC "
            "at 19:00 inside the 16:00 H4 bar) was missed.")
    code(d, "# Bar open t is blocked if some event e satisfies:\n"
            "#   t - before <= e <= t + bar_duration + after\n"
            "lo = self._event_times - bar_td - self.after\n"
            "hi = self._event_times + self.before", lang="python fix")
    code(d, "bars=241 blackout=4 (1.7%)  <-- now correctly flags the FOMC-containing bar",
         lang="output", fg=OK_FG)
    h2(d, "Result: does avoiding news improve returns?")
    table(d, ["Instrument", "Filter", "PF", "Return", "Sharpe"], [
        ["Gold", "OFF / ON", "1.87 / 1.86", "+13.90% / +13.85%", "1.35 / 1.32"],
        ["Silver", "OFF / ON", "1.61 / 1.41", "+11.02% / +7.83%", "0.85 / 0.63"],
        ["AUDUSD", "OFF / ON", "1.66 / 1.67", "+21.13% / +17.70%", "1.35 / 1.32"],
        ["USDJPY", "OFF / ON", "1.32 / 1.32", "+9.11% / +9.11%", "0.64 / 0.64"],
    ])
    para(d, "Honest finding: on price-only backtests the filter is neutral-to-slightly "
            "negative for multi-day swing trades. BUT it was KEPT for live trading, "
            "because backtests model a FIXED spread and cannot see the real spread "
            "blow-outs and slippage during NFP/FOMC. Config: OFF for backtest, ON for live.")

    # ---------------- 12. MT5 CONNECTOR ----------------
    h1(d, "12. Phase 10 — MT5 Connector and Live Bot")
    para(d, "Design principle: separate pure logic (testable anywhere) from MT5 I/O "
            "(Windows only), so the codebase still imports and tests on Linux.")
    code(d, "try:                      # Windows only\n"
            "    import MetaTrader5 as mt5\n"
            "    _MT5_AVAILABLE = True\n"
            "except Exception:\n"
            "    mt5 = None\n"
            "    _MT5_AVAILABLE = False", lang="python — import guard")
    h2(d, "Files created")
    bullets(d, [
        "src/connectors/mt5_connector.py — connect, account_info, symbol_spec (real broker "
        "specs into the risk engine), get_candles, current_tick, positions, "
        "place_market_order with SL/TP, close_position, filling-mode auto-detection, "
        "closed_deals history",
        "src/live/decision.py — PURE decision logic (enter/close/hold/skip/nothing), unit-tested",
        "src/live/portfolio.py — which instrument uses which strategy + fixed params",
        "src/live/trader.py — orchestrator: dry-run, kill switches, news blackout, per-bar loop",
        "src/live/journal.py — CSV journal of actions + equity snapshots",
        "live_trader.py, scripts/check_mt5.py, scripts/dry_run_once.py, scripts/test_order.py, "
        "scripts/report.py",
    ])
    h2(d, "Safety architecture")
    bullets(d, [
        "dry_run mode — computes and logs everything, sends NO orders",
        "Total-drawdown KILL SWITCH (default 20%) — halts the bot entirely",
        "Daily-loss limit (default 6%) — stops new entries for the day",
        "max_open_trades cap",
        "Magic number 990011 — the bot only ever manages its own trades",
        "Acts only on CLOSED bars (drops the still-forming bar) — no repainting",
        "Broker-side SL/TP on every order",
    ])

    # ---------------- 13. WINDOWS ERRORS ----------------
    h1(d, "13. Phase 11 — Windows Setup: every error we hit")
    h2(d, "ERROR — pandas-ta / numba refused to build on Python 3.14")
    error_block(d, "Collecting numba==0.61.2 (from pandas-ta)\n"
                   "  Getting requirements to build wheel ... error\n"
                   "  RuntimeError: Cannot install on Python version 3.14.6;\n"
                   "  only versions >=3.10,<3.14 are supported.\n"
                   "ERROR: Failed to build 'numba' when getting requirements to build wheel")
    para(d, "Resolution: pandas-ta was never actually used — all indicators are "
            "hand-written in pandas. Removed the dependency entirely.")
    code(d, "# requirements.txt after fix\nPyYAML>=6.0\npandas>=2.0\nnumpy>=1.24\npytest>=8.0\n"
            "# MetaTrader5 installed separately on Windows", lang="fix")

    h2(d, "ERROR — scripts not found (wrong branch)")
    error_block(d, "python: can't open file 'C:\\Tradingbot\\scripts\\check_mt5.py':\n"
                   "  [Errno 2] No such file or directory\n"
                   "git pull -> 'Already up to date.'")
    para(d, "Cause: a fresh clone defaults to main, but the work was on the edge-hunting "
            "branch. Resolution: git checkout edge-hunting (later merged to main).")

    h2(d, "ERROR — wrong directory")
    error_block(d, "C:\\Users\\Lingesh K>git pull\n"
                   "fatal: not a git repository (or any of the parent directories): .git")
    para(d, "Cause: a new terminal opens in the home folder. Resolution: cd Tradingbot, "
            "then .venv\\Scripts\\activate.")

    h2(d, "ERROR — Windows backslash escape mangled the command")
    error_block(d, "C:\\Tradingbot>python scripts\\news_dashboard.py\n"
                   "python: can't find '__main__' module in 'C:\\Tradingbot\\scripts'\n"
                   "'ews_dashboard.py' is not recognized as an internal or external command")
    para(d, "Cause: the \\n in scripts\\news_dashboard.py was interpreted as a newline, "
            "splitting the command. Resolution: use forward slashes — "
            "python scripts/news_dashboard.py.")

    h2(d, "ERROR — news feed 404 killed the whole calendar fetch")
    error_block(d, "WARNING trader: Could not refresh news (HTTP Error 404: Not Found);\n"
                   "proceeding without blackout")
    para(d, "Cause: only ff_calendar_thisweek.json exists; lastweek/nextweek 404, and one "
            "failure aborted the loop. Resolution: fetch each week independently.")
    code(d, "for wk in ('lastweek', 'thisweek', 'nextweek'):\n"
            "    try:\n"
            "        events.extend(cal.events(week=wk, min_impact='High'))\n"
            "    except Exception as e:\n"
            "        log.debug(\"news feed '%s' unavailable: %s\", wk, e)", lang="python fix")

    h2(d, "ERROR — git pull blocked by local config changes")
    error_block(d, "error: Your local changes to the following files would be overwritten by merge:\n"
                   "        config/live_config.yaml\n"
                   "Please commit your changes or stash them before you merge. Aborting")
    para(d, "Resolution (short term): git stash; git pull; git stash pop. "
            "Permanent fix in Phase 17: untrack the config and ship a template.")

    h2(d, "Non-fatal warning — venv launcher copy")
    error_block(d, "Unable to copy 'venvlauncher.exe' to 'C:\\Tradingbot\\.venv\\Scripts\\python.exe'")
    para(d, "Harmless — activation, pip installs and all scripts worked normally afterwards.")

    h2(d, "SUCCESS — symbol discovery on the real account")
    para(d, "Two guessed symbol names were wrong and corrected from the diagnostic output:")
    table(d, ["Instrument", "Guess", "Actual (JustMarkets ECN)"], [
        ["Natural Gas", "NGAS.ecn", "XNGUSD.ecn"],
        ["Brent", "UKOIL.ecn", "BRENT.ecn"],
        ["Gold / Silver / Platinum", "—", "XAUUSD.ecn / XAGUSD.ecn / XPTUSD.ecn (correct)"],
        ["GBPJPY / AUDUSD / USDJPY", "—", "GBPJPY.ecn / AUDUSD.ecn / USDJPY.ecn (correct)"],
    ])
    h2(d, "SUCCESS — live order pathway validated on DEMO")
    code(d, "python scripts/test_order.py --confirm", lang="bash")
    code(d, "order result: {'ok': True, 'retcode': 10009, 'order': 2164208849,\n"
            "               'deal': 2024119080, 'price': 4047.16,\n"
            "               'comment': 'Request executed'}\n"
            "ORDER OK. Waiting 3s, then closing it...\n"
            "close result: {'ok': True, 'retcode': 10009, ... 'Request executed'}\n"
            "Done - you should be flat again.", lang="output", fg=OK_FG)
    para(d, "retcode 10009 = TRADE_RETCODE_DONE. Order placement with SL/TP, ECN filling "
            "mode detection, and clean close all confirmed.")

    # ---------------- 14. LOGGING / VPS ----------------
    h1(d, "14. Phase 12 — Performance Logging and VPS 24/7")
    bullets(d, [
        "src/live/journal.py — CSV journal (every action + equity snapshots) to logs/",
        "connector.closed_deals(days) — realized P&L from broker history, filtered by magic",
        "scripts/report.py — live win rate, profit factor, net P&L, per-symbol breakdown, "
        "equity drawdown",
        "docs/VPS_DEPLOYMENT.md — Windows VPS guide",
        "scripts/run_bot.bat — auto-restart launcher (+ Startup-folder autostart)",
    ])
    para(d, "Important discovery documented: MetaTrader 5's built-in 'Virtual Hosting' VPS "
            "runs only MQL5 Expert Advisors, NOT Python. A full Windows VPS is required.")
    code(d, "@echo off\ncd /d \"%~dp0\\..\"\ncall .venv\\Scripts\\activate\n"
            ":loop\necho [%date% %time%] Starting bot...\npython live_trader.py\n"
            "echo Bot exited. Restarting in 30 seconds...\ntimeout /t 30 /nobreak\ngoto loop",
         lang="batch — scripts/run_bot.bat")

    # ---------------- 15. SMC MINING ----------------
    h1(d, "15. Phase 13 — Mining 344 SMC/ICT Indicators")
    para(d, "A second repository (klingesh/EA) contained ~343 TradingView Pine Script "
            "indicators (16 MB) collected from a Telegram channel. These were catalogued "
            "and the distinct concepts translated into Python strategies for testing.")
    h2(d, "Catalogue results")
    table(d, ["Count", "Category"], [
        ["142", "Trend / MA"],
        ["77", "Smart Money Concepts (SMC/ICT)"],
        ["43", "Support / Resistance / Pivots"],
        ["34", "Momentum / Oscillator"],
        ["19", "Divergence / Signals"],
        ["7", "Volume / CVD"],
        ["7", "Sessions / Time"],
        ["7", "Patterns / Harmonic / Fib"],
        ["6", "Supply & Demand"],
    ])
    para(d, "Also found: Pine v6 (229), v5 (112), v4 (1); 10 actual strategy() scripts; "
            "6 groups of identical duplicate files.")
    h2(d, "Concept 1 — Fair Value Gap (FVG)")
    para(d, "Logic confirmed directly from the indicator source (EA/Indi 39, lines 99-104):")
    code(d, "bullFVG   = low > high[2]        // BISI: 3-candle up-gap\n"
            "bearFVG   = high < low[2]        // SIBI: 3-candle down-gap\n"
            "topVal    = bullFVG ? low  : low[2]\n"
            "bottomVal = bullFVG ? high[2] : high", lang="pine script — source logic")
    para(d, "Implemented as src/strategy/fvg.py (retrace into a fresh imbalance, aligned "
            "with the daily trend, one-shot mitigation, ATR stops).")
    para(d, "In-sample on BTC looked great: PF 1.12, +27.3% return. Then walk-forward:")
    table(d, ["Instrument", "PF", "Return", "MaxDD", "Sharpe", "Verdict"], [
        ["BTC", "0.84", "-12.08%", "30.41%", "-0.19", "LOSES (in-sample +27% was overfit)"],
        ["Gold", "1.14", "+12.21%", "26.61%", "0.58", "worse than trend (1.35 / 5.9% DD)"],
        ["Silver", "1.24", "+11.85%", "14.21%", "0.72", "worse than trend"],
        ["AUDUSD", "0.60", "-25.70%", "36.85%", "-0.94", "loses badly"],
        ["USDJPY", "0.65", "-8.94%", "22.79%", "-0.36", "loses"],
    ])
    para(d, "VERDICT: REJECTED. The BTC in-sample +27% collapsing to -12% out-of-sample is "
            "a textbook overfitting trap — caught before any money was risked.", bold=True)
    h2(d, "Concept 2 — Order Blocks (SMC/Supply-Demand)")
    para(d, "src/strategy/order_block.py — the last opposite-colour candle before an "
            "impulsive structural break becomes a demand/supply zone; enter on retest.")
    table(d, ["Instrument", "PF", "Return", "MaxDD", "Sharpe", "Verdict"], [
        ["BTC", "0.71", "-15.46%", "29.08%", "-0.63", "loses"],
        ["Gold", "1.21", "+6.65%", "8.33%", "0.56", "worse than trend"],
        ["Silver", "1.67", "+19.20%", "5.86%", "1.30", "BEATS incumbent trend"],
        ["AUDUSD", "0.36", "-19.02%", "26.78%", "-1.18", "loses badly"],
        ["USDJPY", "1.23", "+9.81%", "9.16%", "0.50", "worse than mean-reversion"],
    ])
    para(d, "One candidate found: Order Block on SILVER is a Pareto improvement over the "
            "incumbent silver trend strategy (+19.2% vs +11%, Sharpe 1.30 vs 0.85, 5.9% vs "
            "6.9% drawdown). Decision: do NOT blind-swap — flagged for demo A/B testing, "
            "because 26 OOS trades is a modest sample and ~15 combinations were tested "
            "(multiple-comparisons risk).")
    h2(d, "Concept 3 — Liquidity Sweep / stop-hunt reversal (ICT)")
    para(d, "src/strategy/liquidity_sweep.py — price pierces a prior N-bar extreme then "
            "closes back inside (trapped traders) = reversal signal.")
    table(d, ["Instrument", "PF", "Return", "Sharpe"], [
        ["BTC", "0.81", "-13.81%", "-0.37"],
        ["Gold", "1.00", "+2.83%", "0.22"],
        ["Silver", "1.16", "+9.82%", "0.53"],
        ["AUDUSD", "1.25", "+9.84%", "0.51"],
        ["USDJPY", "1.21", "+18.26%", "0.54"],
    ])
    para(d, "VERDICT: REJECTED — positive on four instruments but dominated by the "
            "existing strategies everywhere, and loses on BTC.", bold=True)
    h2(d, "Independent research on SMC/ICT")
    para(d, "External sources agree with our own results: there is no verified profitability "
            "for ICT/SMC in academic or tier-1 literature, the widely-quoted '80–90% "
            "order-block fill rates' are unsourced lore, and SMC/ICT largely rebrands "
            "classic support/resistance concepts. Public backtests are mixed-to-negative. "
            "Sources: alphaexcapital.com, phidiaspropfirm.com, lunetrading.com, "
            "opinicusholdings.com (content paraphrased for licensing compliance).")

    # ---------------- 16. VOL TARGETING ----------------
    h1(d, "16. Phase 14 — Volatility Targeting (evidence-based)")
    para(d, "Instead of mining more indicators, techniques with genuine academic support "
            "were researched and added: time-series momentum, cross-sectional momentum, "
            "volatility targeting, and diversification. Sources include Moreira & Muir "
            "'Volatility-Managed Portfolios' (SSRN), vol-scaled momentum studies (MDPI), "
            "and trend-following literature (SSRN/arXiv) — paraphrased for compliance.")
    para(d, "A notable supporting statistic: a study of 66,000 simulated trend trades found "
            "fewer than 7% of trades drove all cumulative profit — precisely the "
            "'low win rate, big winners' skew our own system exhibits.")
    h2(d, "Implementation: src/risk/vol_target.py")
    code(d, "ret     = np.log(close / close.shift(1))\n"
            "rvol    = ret.rolling(lookback).std()\n"
            "target  = rvol.rolling(median_window, min_periods=lookback).median()\n"
            "scalar  = (target / rvol).clip(lower=min_scale, upper=max_scale)\n"
            "scalar  = scalar.shift(1).fillna(1.0)   # no look-ahead", lang="python")
    h2(d, "Result — NOT universally good")
    table(d, ["Instrument", "Strategy", "Sharpe off->on", "Return off->on", "MaxDD off->on"], [
        ["BTC", "trend", "0.76 -> 0.31", "+19.6% -> +6.2%", "13.5% -> 15.5%"],
        ["Gold", "trend", "1.35 -> 0.88", "+13.9% -> +8.9%", "5.9% -> 8.3%"],
        ["Silver", "trend", "0.85 -> 0.99", "+11.0% -> +14.7%", "6.9% -> 6.2%"],
        ["AUDUSD", "mean-rev", "1.27 -> 1.54", "+18.9% -> +25.0%", "9.1% -> 8.8%"],
        ["USDJPY", "mean-rev", "1.18 -> 1.54", "+17.9% -> +24.7%", "11.9% -> 8.4%"],
    ])
    para(d, "It HURTS trend-following (de-risking during volatile expansions cuts the "
            "skew-driven tail winners trend depends on) but clearly HELPS mean-reversion "
            "(which is endangered by high volatility). Decision: apply SELECTIVELY — ON "
            "for the mean-reversion sleeve, OFF for the trend sleeve. A mechanism-based "
            "rule, not per-instrument curve fitting.", bold=True)
    h2(d, "Test that caught a subtlety")
    error_block(d, "FAILED tests/test_vol_target.py::test_derisks_in_high_vol\n"
                   "assert np.mean(s[-50:]) < 1.0")
    para(d, "Cause: with a short median window, a persistent high-vol regime lets the "
            "median catch up, so de-risking stops — correct regime-relative behaviour. "
            "The test was corrected to use a longer calibration window.")

    # ---------------- 17. UNIVERSE ----------------
    h1(d, "17. Phase 15 — Universe Expansion & Diversification")
    para(d, "Expanded from 5 to 25 tested instruments: 11 FX (majors + JPY/EUR crosses), "
            "7 commodities, 4 indices, 3 crypto. Nikkei was dropped (only 968 H4 bars).")
    code(d, "python experiments/universe_scan.py", lang="bash")
    para(d, "Screening rule (a 'keeper' must clear all): >=10 OOS trades, PF >= 1.2, "
            "Sharpe >= 0.4, positive return. Both strategy types tested on every "
            "instrument; the better one kept.")
    h2(d, "9 keepers from 25 instruments")
    table(d, ["Instrument", "Class", "Strategy", "Sharpe", "Return", "MaxDD"], [
        ["XAUUSD", "commodity", "trend", "2.16", "+36.4%", "10.7%"],
        ["AUDUSD", "forex", "mean-rev", "1.63", "+29.7%", "8.9%"],
        ["XAGUSD", "commodity", "trend", "1.42", "+19.9%", "6.9%"],
        ["NATGAS", "commodity", "trend", "1.20", "+22.2%", "6.4%"],
        ["BTC", "crypto", "trend", "0.92", "+24.6%", "10.1%"],
        ["PLATINUM", "commodity", "trend", "0.74", "+7.9%", "6.7%"],
        ["BRENT", "commodity", "mean-rev", "0.73", "+11.2%", "9.1%"],
        ["GBPJPY", "forex", "trend", "0.67", "+12.8%", "11.5%"],
        ["USDJPY", "forex", "mean-rev", "0.64", "+8.4%", "11.4%"],
    ])
    para(d, "Rejected: EURUSD, GBPUSD, USDCAD, USDCHF, NZDUSD, EURJPY, EURGBP, AUDJPY, "
            "WTI, COPPER, and ALL indices (S&P 0.36, Nasdaq 0.44, Dow -0.38, DAX 0.20), "
            "plus ETH (0.38) and SOL (-0.12).")
    h2(d, "The diversification result")
    code(d, "=== BLENDED PORTFOLIO (equal-weight keepers) ===\n"
            "  Instruments blended ... 9\n"
            "  Avg individual Sharpe . 1.02\n"
            "  BLENDED Sharpe ........ 2.28   <- diversification benefit\n"
            "  Blended max drawdown .. 1.96%\n"
            "  Blended total return .. +19.66%", lang="output", fg=OK_FG)
    para(d, "Combining uncorrelated edges more than doubled the Sharpe and collapsed "
            "drawdown. IMPORTANT CAVEAT: the exact figure is optimistic — keepers were "
            "selected (bias), 5 of 9 are correlated commodities, the window is short, and "
            "flat/non-trading days dilute measured volatility. The principle is solid; "
            "expect a real but smaller benefit live.")

    # ---------------- 18. NEWS TOOLS ----------------
    h1(d, "18. Phase 16 — News Awareness Tools")
    para(d, "Request: have the bot read forex news, economic calendars and geopolitical "
            "news (e.g. Trump speeches, Strait of Hormuz) and trade on positive/negative "
            "sentiment.")
    h2(d, "Why news-sentiment AUTO-TRADING was deliberately NOT built")
    bullets(d, [
        "LATENCY: news moves gold/oil in milliseconds; institutions and HFT react first, "
        "and spreads blow out 5–20x. A retail bot is structurally too slow.",
        "SENTIMENT != DIRECTION: markets are forward-looking ('buy the rumour, sell the "
        "news'). A good CPI print can send gold up OR down depending on Fed implications.",
        "UNBACKTESTABLE: there is no clean historical archive of geopolitical headlines "
        "with timestamps, sentiment and market reactions — so no edge can be validated. "
        "This violates the discipline used for every other component.",
    ])
    h2(d, "What was built instead")
    bullets(d, [
        "scripts/calendar_report.py — this week's high-impact events + per-instrument "
        "blackout schedule",
        "src/news/sentiment.py — Google News RSS fetch + finance-tuned tone lexicon "
        "(score in [-1, +1])",
        "scripts/news_dashboard.py — per-instrument and macro/geopolitical sentiment radar",
    ])
    para(d, "Design: news makes the HUMAN smarter (dashboard) and protects the BOT "
            "(blackout filter), but never fires trades automatically.", bold=True)
    h2(d, "Live demonstration output")
    code(d, "--- MACRO / GEOPOLITICAL ---\n"
            "  Strait of Hormuz     -0.40 NEGATIVE tone  (0+/6-/9o of 15)\n"
            "  Fed / rates          +0.00 neutral/mixed  (2+/2-/11o of 15)\n"
            "  US politics          -0.13 neutral/mixed  (0+/2-/13o of 15)\n"
            "  Middle East / oil    -0.36 NEGATIVE tone  (2+/8-/5o of 15)\n"
            "--- PORTFOLIO INSTRUMENTS ---\n"
            "  GOLD                 +0.29 POSITIVE tone  (7+/2-/6o of 15)\n"
            "  PLATINUM             +0.40 POSITIVE tone  (6+/0-/9o of 15)\n"
            "  USDJPY               -0.27 NEGATIVE tone  (2+/6-/7o of 15)\n"
            "  BRENT                -0.20 NEGATIVE tone  (3+/6-/6o of 15)", lang="output")

    # ---------------- 19. CONSOLIDATION ----------------
    h1(d, "19. Phase 17 — Consolidation and Config Decoupling")
    para(d, "Merged the edge-hunting branch into main (fast-forward + one merge to "
            "reconcile a dry_run change made on GitHub), then aligned both branches so "
            "future clones just use main.")
    h2(d, "Permanent fix for the config/pull conflict")
    bullets(d, [
        "git rm --cached config/live_config.yaml  (untrack, keep the file locally)",
        "Added config/live_config.yaml to .gitignore",
        "Created config/live_config.example.yaml as the tracked template",
        "Documented the one-time copy step in docs/WINDOWS_SETUP.md",
    ])
    para(d, "Result: personal settings (symbols, risk, dry_run) never conflict with "
            "git pull again.")

    # ---------------- 20. FILE INVENTORY ----------------
    h1(d, "20. Complete File Inventory")
    h2(d, "Source code — src/")
    table(d, ["File", "Language", "Purpose"], [
        ["src/risk/position_sizing.py", "Python", "Lot sizing from risk %, hard caps, min-lot policy"],
        ["src/risk/vol_target.py", "Python", "Volatility-target risk scalar (regime-relative)"],
        ["src/strategy/base.py", "Python", "Abstract Strategy interface"],
        ["src/strategy/indicators.py", "Python", "EMA, RSI (Wilder), ATR"],
        ["src/strategy/ema_rsi_swing.py", "Python", "Trend-following + HTF filter"],
        ["src/strategy/mean_reversion.py", "Python", "Bollinger + RSI mean-reversion"],
        ["src/strategy/breakout.py", "Python", "Donchian channel breakout"],
        ["src/strategy/fvg.py", "Python", "Fair Value Gap continuation (rejected)"],
        ["src/strategy/order_block.py", "Python", "Order Block continuation (silver candidate)"],
        ["src/strategy/liquidity_sweep.py", "Python", "ICT stop-hunt reversal (rejected)"],
        ["src/backtest/engine.py", "Python", "Event-driven backtester (NumPy hot loop)"],
        ["src/backtest/metrics.py", "Python", "Expectancy, PF, Sharpe, drawdown, R-multiples"],
        ["src/backtest/walk_forward.py", "Python", "Rolling IS/OOS validation harness"],
        ["src/data/loader.py", "Python", "OKX crypto OHLCV + CSV cache"],
        ["src/data/yahoo_loader.py", "Python", "Yahoo FX/commodity/index daily + H4 resample"],
        ["src/news/calendar.py", "Python", "Deterministic US calendar + ForexFactory feed"],
        ["src/news/filter.py", "Python", "Blackout mask (bar-duration aware)"],
        ["src/news/sentiment.py", "Python", "Headline fetch + finance tone lexicon"],
        ["src/connectors/mt5_connector.py", "Python", "MetaTrader 5 wrapper (Windows)"],
        ["src/live/decision.py", "Python", "Pure decision logic (unit-tested)"],
        ["src/live/portfolio.py", "Python", "Instrument -> strategy mapping + params"],
        ["src/live/trader.py", "Python", "Live orchestrator + kill switches"],
        ["src/live/journal.py", "Python", "CSV action/equity journal"],
    ])
    h2(d, "Scripts, config, tests, docs")
    table(d, ["File", "Type", "Purpose"], [
        ["live_trader.py", "Python", "Entry point for continuous live/demo trading"],
        ["backtest_run.py", "Python", "Single backtest runner"],
        ["scripts/check_mt5.py", "Python", "Connection test + symbol/spec discovery"],
        ["scripts/dry_run_once.py", "Python", "One-shot evaluation of all instruments"],
        ["scripts/test_order.py", "Python", "Validate order pathway on demo (0.01 lot)"],
        ["scripts/report.py", "Python", "Live performance report from broker history"],
        ["scripts/calendar_report.py", "Python", "Weekly events + blackout schedule"],
        ["scripts/news_dashboard.py", "Python", "Sentiment radar (informational)"],
        ["scripts/run_bot.bat", "Batch", "Auto-restart launcher for the VPS"],
        ["config/live_config.example.yaml", "YAML", "Config template (gitignored real copy)"],
        ["config/config.yaml", "YAML", "Original scaffold risk config"],
        ["requirements.txt", "Text", "Python dependencies"],
        ["tests/test_position_sizing.py", "Python", "8 risk-engine tests"],
        ["tests/test_news_filter.py", "Python", "5 calendar/blackout tests"],
        ["tests/test_decision.py", "Python", "8 live-decision tests"],
        ["tests/test_vol_target.py", "Python", "4 volatility-target tests"],
        ["tests/test_sentiment.py", "Python", "6 sentiment-scoring tests"],
        ["experiments/ (10 files)", "Python", "All research experiments, reproducible"],
        ["docs/ (7 files)", "Markdown", "Research log, setup, VPS, news tools, reports"],
    ])

    # ---------------- 21. STRATEGIES ----------------
    h1(d, "21. All Strategies: Logic and Verdicts")
    h2(d, "1. EMA + RSI Trend-Following (SELECTED)")
    code(d, "fast_above  = EMA(close, fast) > EMA(close, slow)\n"
            "cross_up    = fast_above AND NOT fast_above.shift(1)\n"
            "cross_down  = NOT fast_above AND fast_above.shift(1)\n"
            "long_entry  = cross_up   AND RSI > 50  AND htf_dir == +1\n"
            "short_entry = cross_down AND RSI < 50  AND htf_dir == -1\n"
            "stop_loss   = entry -/+ sl_atr_mult * ATR\n"
            "take_profit = entry +/- tp_rr * (stop distance)", lang="pseudocode")
    para(d, "Used on: Gold, Silver, NatGas, Platinum, GBPJPY (+ optional BTC). "
            "Vol-targeting OFF.")
    h2(d, "2. Bollinger + RSI Mean-Reversion (SELECTED)")
    code(d, "mid   = SMA(close, period);  std = STDEV(close, period)\n"
            "upper = mid + num_std * std;  lower = mid - num_std * std\n"
            "long_entry  = close < lower AND RSI < rsi_oversold\n"
            "short_entry = close > upper AND RSI > rsi_overbought\n"
            "stops: ATR-based; small tp_rr (revert toward the mean)", lang="pseudocode")
    para(d, "Used on: AUDUSD, USDJPY, Brent. Vol-targeting ON.")
    h2(d, "3. Donchian Breakout (tested, not selected)")
    code(d, "upper = high.rolling(channel).max().shift(1)\n"
            "lower = low.rolling(channel).min().shift(1)\n"
            "broke_up   = close > upper AND close.shift(1) <= upper.shift(1)\n"
            "broke_down = close < lower AND close.shift(1) >= lower.shift(1)", lang="pseudocode")
    para(d, "Highest raw return on BTC (+26.7%) but 28.9% drawdown — poor risk-adjusted "
            "profile versus the trend strategy.")
    h2(d, "4. Fair Value Gap (REJECTED)")
    code(d, "bullish FVG: low[0] > high[2]   -> zone [high[2], low[0]]  (demand)\n"
            "bearish FVG: high[0] < low[2]   -> zone [high[0], low[2]]  (supply)\n"
            "entry: price retraces into a fresh zone, aligned with daily trend\n"
            "zone invalidated by a close through it, or after max_zone_age bars",
         lang="pseudocode")
    h2(d, "5. Order Block (REJECTED except a silver candidate)")
    code(d, "bullish OB: previous candle bearish AND close > previous high\n"
            "            AND (close - open) >= impulse_atr * ATR\n"
            "            -> demand zone = [low[-1], high[-1]]\n"
            "entry: retest of the zone in the trend direction", lang="pseudocode")
    h2(d, "6. Liquidity Sweep (REJECTED)")
    code(d, "prior_low  = low.rolling(lookback).min().shift(1)\n"
            "prior_high = high.rolling(lookback).max().shift(1)\n"
            "bull_sweep = low  < prior_low  AND close > prior_low\n"
            "bear_sweep = high > prior_high AND close < prior_high", lang="pseudocode")

    # ---------------- 22. ERROR LOG ----------------
    h1(d, "22. Complete Error Log and Resolutions")
    table(d, ["#", "Error", "Cause", "Resolution"], [
        ["1", "HTTPError 451 (Binance)", "Geo-restricted API", "Switched to OKX"],
        ["2", "HTTPError 401 (CryptoCompare, FMP)", "Auth required", "Used free no-auth sources"],
        ["3", "HTTPError 403 / 400 (Bybit, Coinbase)", "Blocked / bad params", "Used OKX + Yahoo"],
        ["4", "Yahoo returned monthly bars", "range=max downsamples", "Explicit period1/period2"],
        ["5", "HTTPError 422 (Yahoo hourly)", "Range > ~729 days", "729d of 1h, resample to H4"],
        ["6", "Stooq bot-blocked (HTML)", "Anti-scraping", "Used Yahoo instead"],
        ["7", "Backtest timeout (120s)", "Slow pandas .iloc loop", "NumPy hot loop (~20x faster)"],
        ["8", "Misleading 'POSITIVE EDGE' label", "Threshold too loose", "PF>1.05 AND expR>0.03"],
        ["9", "TypeError naive vs aware datetime", "Mixed tz in numpy compare", "Normalize to naive-UTC"],
        ["10", "Blackout missed mid-bar events", "Only bar open checked", "Bar-duration-aware mask"],
        ["11", "numba/pandas-ta build failure (Py3.14)", "Unsupported Python", "Removed unused dependency"],
        ["12", "MetaTrader5 not installed", "Fresh venv", "pip install MetaTrader5 (Windows)"],
        ["13", "News 404 aborted calendar fetch", "lastweek/nextweek absent", "Fetch each week independently"],
        ["14", "Wrong symbol guesses (NGAS/UKOIL)", "Broker-specific naming", "check_mt5.py discovery"],
        ["15", "'not a git repository'", "Terminal in home folder", "cd Tradingbot first"],
        ["16", "Scripts not found on VPS", "Fresh clone on main branch", "checkout branch, later merged"],
        ["17", "Backslash mangled command", "\\n read as newline", "Use forward slashes"],
        ["18", "git pull blocked by config edits", "Tracked config file", "Untrack + gitignore + template"],
        ["19", "venvlauncher.exe copy warning", "Windows quirk", "Harmless; venv worked"],
        ["20", "Push rejected (fetch first)", "Remote main had merges", "Merged origin/main, then pushed"],
        ["21", "test_derisks_in_high_vol failed", "Short median window catches up", "Longer calibration window"],
    ])

    # ---------------- 23. FINAL ----------------
    h1(d, "23. Final Portfolio and Results")
    table(d, ["Instrument", "Broker symbol", "Strategy", "Vol-target"], [
        ["Gold", "XAUUSD.ecn", "H4 trend (EMA+RSI + daily filter)", "OFF"],
        ["Silver", "XAGUSD.ecn", "H4 trend", "OFF"],
        ["Natural Gas", "XNGUSD.ecn", "H4 trend", "OFF"],
        ["Platinum", "XPTUSD.ecn", "H4 trend", "OFF"],
        ["GBPJPY", "GBPJPY.ecn", "H4 trend", "OFF"],
        ["AUDUSD", "AUDUSD.ecn", "H4 mean-reversion", "ON"],
        ["USDJPY", "USDJPY.ecn", "H4 mean-reversion", "ON"],
        ["Brent", "BRENT.ecn", "H4 mean-reversion", "ON"],
        ["BTC (optional)", "BTCUSD.ecn", "H4 trend", "OFF"],
    ])
    h2(d, "Risk configuration in production")
    code(d, "dry_run: false                      # armed on DEMO\n"
            "risk_percent_per_trade: 2.0\n"
            "max_risk_percent_per_trade: 5.0\n"
            "max_open_trades: 8\n"
            "max_daily_loss_percent: 6.0         # halt new entries for the day\n"
            "max_total_drawdown_percent: 20.0    # KILL SWITCH\n"
            "news_filter: enabled true, +/-120 minutes\n"
            "magic_number: 990011", lang="yaml")
    h2(d, "Verified live on the demo account")
    code(d, "Connected. Login 1100219238  Balance 10494.94 USD  [DRY RUN]\n"
            "Refreshed news calendar: 11 high-impact events this week\n"
            "[GOLD/XAUUSD.ecn]     bar=2026-07-24 16:00 close=4067.11  signal=+0  volscale=1.00\n"
            "[SILVER/XAGUSD.ecn]   bar=2026-07-24 16:00 close=58.63    signal=+0  volscale=1.00\n"
            "[NATGAS/XNGUSD.ecn]   bar=2026-07-24 16:00 close=2.915    signal=+0  volscale=1.00\n"
            "[PLATINUM/XPTUSD.ecn] bar=2026-07-24 16:00 close=1598.97  signal=+0  volscale=1.00\n"
            "[GBPJPY/GBPJPY.ecn]   bar=2026-07-24 16:00 close=218.362  signal=+0  volscale=1.00\n"
            "[AUDUSD/AUDUSD.ecn]   bar=2026-07-24 16:00 close=0.69905  signal=+0  volscale=0.95\n"
            "[USDJPY/USDJPY.ecn]   bar=2026-07-24 16:00 close=163.781  signal=+0  volscale=0.92\n"
            "[BRENT/BRENT.ecn]     bar=2026-07-24 16:00 close=91.02    signal=+0  volscale=1.16\n"
            "One-shot check complete.", lang="output", fg=OK_FG)
    para(d, "Note the volscale values: 1.00 for trend slots (targeting off) and 0.92–1.16 "
            "for mean-reversion slots (targeting active) — exactly as designed.")
    para(d, "Test suite: 31 automated tests passing.", bold=True)

    # ---------------- 24. LESSONS ----------------
    h1(d, "24. Key Lessons Learned")
    bullets(d, [
        "Win rate is a vanity metric. Expectancy and risk/reward determine profitability. "
        "Our profitable strategies win only ~40–57% of trades; the highest-win-rate "
        "strategy tested was the worst performer.",
        "Backtest everything; trust only out-of-sample. The FVG strategy showed +27% "
        "in-sample on BTC and -12% out-of-sample. Walk-forward caught it before any risk.",
        "Match the strategy to the market's character. USDJPY: -8% with trend, +18% with "
        "mean-reversion. Commodities trend; FX majors mean-revert.",
        "Edges come from mechanisms, not secret indicators. 344 SMC/ICT indicators yielded "
        "one marginal candidate; evidence-based methods (trend, vol targeting, "
        "diversification) delivered the real gains.",
        "Risk management is the product. Sizing, stops, caps, kill switches and the news "
        "blackout matter more than any entry signal.",
        "Diversification is the closest thing to a free lunch. Blending uncorrelated edges "
        "roughly doubled the risk-adjusted return and slashed drawdown.",
        "Don't automate what you can't validate. News-sentiment auto-trading was refused "
        "on principle: unbacktestable and latency-disadvantaged.",
        "Enhancements can be regime-specific. Volatility targeting helps mean-reversion and "
        "hurts trend-following — the average alone would have hidden this.",
        "Separate pure logic from I/O. It let every rule be unit-tested on Linux while the "
        "MT5 connector ran only on Windows.",
    ])

    # ---------------- 25. MANUAL ----------------
    h1(d, "25. Operating Manual (commands reference)")
    h2(d, "First-time setup (Windows / VPS)")
    code(d, "git clone https://github.com/klingesh/Tradingbot.git\n"
            "cd Tradingbot\n"
            "python -m venv .venv\n"
            ".venv\\Scripts\\activate\n"
            "pip install -r requirements.txt\n"
            "pip install MetaTrader5\n"
            "copy config\\live_config.example.yaml config\\live_config.yaml", lang="batch")
    h2(d, "Verify and run")
    code(d, "python scripts/check_mt5.py        # connection + symbol/spec discovery\n"
            "python scripts/dry_run_once.py     # one-shot evaluation of all instruments\n"
            "python scripts/test_order.py --confirm   # validate order pathway (DEMO!)\n"
            "python live_trader.py              # continuous loop\n"
            "scripts\\run_bot.bat               # continuous + auto-restart (VPS)", lang="batch")
    h2(d, "Monitor")
    code(d, "python scripts/report.py           # live win rate, PF, net P&L, per-symbol\n"
            "python scripts/calendar_report.py  # week's events + blackout schedule\n"
            "python scripts/news_dashboard.py   # sentiment radar", lang="batch")
    h2(d, "Research / re-validate")
    code(d, "python -m pytest tests/ -q                      # 31 tests\n"
            "python experiments/universe_scan.py             # scan all instruments\n"
            "python experiments/walk_forward_test.py         # OOS validation\n"
            "python experiments/vol_target_test.py           # vol-targeting impact\n"
            "python experiments/smc_concepts_test.py         # OB + sweep tests\n"
            "python experiments/fvg_test.py                  # FVG test", lang="batch")
    para(d, "MT5 prerequisites: Tools -> Options -> Expert Advisors -> Allow algorithmic "
            "trading, the toolbar 'Algo Trading' button green, instruments added to Market "
            "Watch, and the terminal logged into the DEMO account.")

    d.add_paragraph()
    para(d, "End of documentation. Status: deployed 24/7 on a Windows VPS, trading a "
            "JustMarkets DEMO account, pending 1–3 months of forward testing before any "
            "consideration of real capital.", italic=True, size=10)

    d.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    build()
